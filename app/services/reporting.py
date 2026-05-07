from __future__ import annotations

from datetime import date, datetime, timezone
from decimal import Decimal
from pathlib import Path
from uuid import uuid4

from docx import Document
from jinja2.sandbox import SandboxedEnvironment
from sqlalchemy.orm import Session

from app.models import (
    Account,
    AccountType,
    EntryType,
    Event,
    GeneratedReport,
    PlaceholderDefinition,
    Project,
    ReportTemplate,
    TemplateFormat,
    TemplateType,
    Transaction,
    TransactionLine,
)

_safe_env = SandboxedEnvironment(autoescape=False)
REPORT_OUTPUT_DIR = Path("generated_reports")
REPORT_OUTPUT_DIR.mkdir(exist_ok=True)
TEMPLATE_UPLOAD_DIR = Path("report_templates")
TEMPLATE_UPLOAD_DIR.mkdir(exist_ok=True)


def render_template(template_html: str, context: dict) -> str:
    return _safe_env.from_string(template_html).render(**context)


def _base_line_query(db: Session):
    return (
        db.query(TransactionLine)
        .join(Account, Account.id == TransactionLine.account_id)
        .join(Transaction, Transaction.id == TransactionLine.transaction_id)
    )


def _apply_filters(q, filters: dict):
    if filters.get("project_id"):
        q = q.filter(TransactionLine.project_id == int(filters["project_id"]))
    if filters.get("budget_head_id"):
        q = q.filter(TransactionLine.budget_head_id == int(filters["budget_head_id"]))
    if filters.get("from_date"):
        q = q.filter(Transaction.txn_date >= datetime.fromisoformat(str(filters["from_date"])))
    if filters.get("to_date"):
        q = q.filter(Transaction.txn_date <= datetime.fromisoformat(str(filters["to_date"])))
    if filters.get("entry_type"):
        q = q.filter(TransactionLine.entry_type == EntryType(filters["entry_type"]))
    if filters.get("account_type"):
        q = q.filter(Account.account_type == AccountType(filters["account_type"]))
    return q


def _metric_total(db: Session, filters: dict | None = None) -> Decimal:
    q = _base_line_query(db)
    q = _apply_filters(q, filters or {})
    total = Decimal("0")
    for line in q.all():
        total += Decimal(line.amount)
    return total


def evaluate_expression(db: Session, expression: dict, inherited_filters: dict | None = None) -> Decimal:
    inherited_filters = inherited_filters or {}
    op = expression.get("op", "metric")

    if op == "metric":
        metric_name = expression.get("metric", "total")
        filters = {**inherited_filters, **(expression.get("filters") or {})}
        if metric_name == "expense_total":
            filters.setdefault("entry_type", EntryType.debit.value)
            filters.setdefault("account_type", AccountType.expense.value)
            return _metric_total(db, filters)
        if metric_name == "income_total":
            filters.setdefault("entry_type", EntryType.credit.value)
            filters.setdefault("account_type", AccountType.income.value)
            return _metric_total(db, filters)
        if metric_name == "total":
            return _metric_total(db, filters)
        raise ValueError(f"Unsupported metric: {metric_name}")

    operands = [evaluate_expression(db, child, inherited_filters) for child in (expression.get("operands") or [])]
    if op == "add":
        return sum(operands, Decimal("0"))
    if op == "subtract":
        if not operands:
            return Decimal("0")
        head, *tail = operands
        return head - sum(tail, Decimal("0"))
    if op == "multiply":
        result = Decimal("1")
        for v in operands:
            result *= v
        return result
    if op == "divide":
        if len(operands) != 2:
            raise ValueError("Divide requires exactly two operands")
        if operands[1] == 0:
            return Decimal("0")
        return operands[0] / operands[1]

    raise ValueError(f"Unsupported op: {op}")


def _format_placeholder_key(name: str) -> str:
    return f"<<{name}>>"


def build_uc_context(db: Session, project_id: int | None, from_date: date | None, to_date: date | None):
    filters = {}
    if project_id:
        filters["project_id"] = project_id
    if from_date:
        filters["from_date"] = f"{from_date}T00:00:00"
    if to_date:
        filters["to_date"] = f"{to_date}T23:59:59"

    expense_total = evaluate_expression(db, {"op": "metric", "metric": "expense_total"}, filters)
    income_total = evaluate_expression(db, {"op": "metric", "metric": "income_total"}, filters)
    net_total = expense_total - income_total

    project_name = "Centre"
    if project_id:
        project = db.get(Project, project_id)
        if project:
            project_name = project.name

    context = {
        "project_name": project_name,
        "from_date": from_date,
        "to_date": to_date,
        "total_spent": round(float(expense_total), 2),
        "total_expenses": round(float(expense_total), 2),
        "total_incomes": round(float(income_total), 2),
        "net_utilization": round(float(net_total), 2),
        _format_placeholder_key("ProjectName"): project_name,
        _format_placeholder_key("FromDate"): str(from_date or ""),
        _format_placeholder_key("ToDate"): str(to_date or ""),
        _format_placeholder_key("TotalExpenses"): str(round(expense_total, 2)),
        _format_placeholder_key("TotalIncome"): str(round(income_total, 2)),
        _format_placeholder_key("NetUtilization"): str(round(net_total, 2)),
    }

    defs = db.query(PlaceholderDefinition).all()
    for item in defs:
        val = evaluate_expression(db, item.expression_json or {}, item.filters_json or {})
        context[item.name] = str(round(val, 2))
        context[_format_placeholder_key(item.name)] = str(round(val, 2))

    return context


def build_newsletter_context(db: Session, from_date: date | None, to_date: date | None):
    q = db.query(Event)
    if from_date:
        q = q.filter(Event.start_at >= from_date)
    if to_date:
        q = q.filter(Event.start_at <= to_date)
    events = q.order_by(Event.start_at.desc()).all()
    return {
        "from_date": from_date,
        "to_date": to_date,
        "events": events,
    }


def _replace_in_docx(src_path: Path, dest_path: Path, mapping: dict[str, str]) -> None:
    doc = Document(str(src_path))

    def replace_text(text: str) -> str:
        out = text
        for key, value in mapping.items():
            out = out.replace(key, str(value))
            out = out.replace(f"{{{{ {key} }}}}", str(value))
            out = out.replace(f"{{{{{key}}}}}", str(value))
        return out

    for paragraph in doc.paragraphs:
        paragraph.text = replace_text(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = replace_text(cell.text)

    doc.save(str(dest_path))


def _replace_in_doc_binary(src_path: Path, dest_path: Path, mapping: dict[str, str]) -> None:
    content = src_path.read_bytes()
    text = content.decode("latin-1", errors="ignore")
    for key, value in mapping.items():
        text = text.replace(key, str(value))
    dest_path.write_bytes(text.encode("latin-1", errors="ignore"))


def _create_word_report(template: ReportTemplate, context: dict) -> tuple[str, str]:
    if not template.template_file_path:
        raise ValueError("Template file path missing for word template")

    src = Path(template.template_file_path)
    if not src.exists():
        raise ValueError("Template file not found on server")

    ext = ".docx" if template.template_format == TemplateFormat.docx else ".doc"
    out_path = REPORT_OUTPUT_DIR / f"report_{uuid4().hex}{ext}"

    if template.template_format == TemplateFormat.docx:
        _replace_in_docx(src, out_path, context)
        return ("", str(out_path))

    _replace_in_doc_binary(src, out_path, context)
    return ("", str(out_path))


def generate_report(
    db: Session,
    template_id: int,
    project_id: int | None,
    from_date: date | None,
    to_date: date | None,
    parameters: dict,
    user_id: int,
):
    template = db.get(ReportTemplate, template_id)
    if not template:
        raise ValueError("Template not found")

    if template.type in (TemplateType.uc, TemplateType.soe):
        context = build_uc_context(db, project_id, from_date, to_date)
    elif template.type == TemplateType.newsletter:
        context = build_newsletter_context(db, from_date, to_date)
    else:
        context = {}

    context.update(parameters or {})

    if template.template_format == TemplateFormat.html:
        html = render_template(template.html_template, context)
        report_file = REPORT_OUTPUT_DIR / f"report_{uuid4().hex}.html"
        report_file.write_text(html, encoding="utf-8")
        html_output = html
        file_path = str(report_file)
        output_format = "HTML"
    else:
        html_output, file_path = _create_word_report(template, context)
        output_format = template.template_format.value

    report = GeneratedReport(
        template_id=template_id,
        project_id=project_id,
        from_date=from_date,
        to_date=to_date,
        parameters=parameters,
        html_output=html_output,
        pdf_path=file_path,
        output_format=output_format,
        created_by_id=user_id,
        updated_by_id=user_id,
    )
    db.add(report)
    db.commit()
    db.refresh(report)
    return report
